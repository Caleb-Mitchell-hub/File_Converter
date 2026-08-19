"""Word 转换器：实现 docx/doc <-> pdf 的双向转换。

支持的转换组合：

    - ``.docx`` -> ``.pdf``
    - ``.doc``  -> ``.pdf``
    - ``.pdf``  -> ``.docx``

引擎选择策略（仅 Word -> PDF 方向）：

    1. 检测到本机有 Microsoft Office（Windows/macOS）时，优先使用
       :mod:`docx2pdf`（其内部封装 pywin32 / Word for Mac Automation）。
    2. 未检测到 MS Office 但存在 LibreOffice 时，回退到
       ``soffice --headless --convert-to pdf``。
    3. 都没有时，抛 :class:`ConversionError`。

注意：Linux 上 :mod:`docx2pdf` 不可用，因此会自动落到 LibreOffice。

PDF -> DOCX 方向采用 ``pdfplumber`` 抽取文本，再用 ``python-docx`` 写回
Word。该方案简单稳定，但 **会丢失原 PDF 的版式信息**（字体、表格、列、
页眉页脚、矢量图等），每个 PDF 页面会按顺序写入 Word 段落，页与页之间
插入分页符。如需高保真还原，请使用专业 OCR / PDF 反向工程工具。
"""

from __future__ import annotations

import re
import shutil
import subprocess
from pathlib import Path
from typing import Callable, ClassVar, Optional, Tuple

from ..core.base import BaseConverter, ConversionError, PathLike
from ..core.logger import get_logger
from ..utils.paths import ensure_dir, unique_output_path
from ..utils.platform import office_status


# ``docx2pdf`` 在 Linux 上 import 即抛 ImportError，必须做保护。
try:  # pragma: no cover - 由运行环境决定是否可用
    from docx2pdf import convert as _docx2pdf_convert  # type: ignore
except Exception:  # noqa: BLE001 - 任何导入错误都视作不可用
    _docx2pdf_convert = None  # type: ignore[assignment]


# ``pywin32`` 仅 Windows 才有；macOS 上也需要 ``pyobjc`` 才能用 win32com。
# 这里统一保护，非 Windows 直接置空即可。注意 ``win32com`` 本身可能
# 没有被注册到 ``sys.modules``（导入失败时），所以仅保存 client 对象。
_win32_client = None  # type: ignore[var-annotated]
try:  # pragma: no cover - 平台相关
    import win32com.client as _win32_client  # type: ignore
    _HAS_WIN32 = True
except Exception:  # noqa: BLE001
    _HAS_WIN32 = False


# PDF -> DOCX 依赖。
try:  # pragma: no cover - 由运行环境决定
    import pdfplumber  # type: ignore
except Exception:  # noqa: BLE001
    pdfplumber = None  # type: ignore[assignment]

try:  # pragma: no cover - 由运行环境决定
    from docx import Document as _DocxDocument  # type: ignore
except Exception:  # noqa: BLE001
    _DocxDocument = None  # type: ignore[assignment]


# wdFormatPDF = 17，是 MS Word 常量 ``Word.WdSaveFormat.wdFormatPDF``。
_WD_FORMAT_PDF = 17

# 子进程默认超时（秒），避免 LibreOffice 在异常情况下挂死主进程。
_LIBREOFFICE_TIMEOUT = 180


class WordConverter(BaseConverter):
    """Word 文档（doc/docx）与 PDF 之间的转换器。

    Word -> PDF 方向会调用本机 Office / LibreOffice 渲染引擎；PDF -> Word
    方向则通过 ``pdfplumber`` 抽文本，再用 ``python-docx`` 重新拼装
    Word 文档（保真度有限，详见模块级 docstring）。

    Attributes:
        name: 转换器名。
        supported_pairs: 支持的 (源扩展名, 目标扩展名) 组合。
    """

    name: ClassVar[str] = "WordConverter"
    supported_pairs: ClassVar[Tuple[Tuple[str, str], ...]] = (
        (".docx", ".pdf"),
        (".doc", ".pdf"),
        (".pdf", ".docx"),
    )

    def __init__(self, *, engine: Optional[str] = None, **kwargs: object) -> None:
        """初始化。

        Args:
            engine: 强制指定 Word -> PDF 的渲染引擎。取值之一：

                - ``"docx2pdf"``：使用 docx2pdf（需本机有 MS Office）。
                - ``"libreoffice"``：使用 ``soffice --headless``。
                - ``"pywin32"``：Windows 下用 win32com 直接调用 MS Word。
                - ``None``：自动选择（默认）。

            **kwargs: 预留扩展参数（当前未使用，保留以备将来扩展）。
        """
        super().__init__(**kwargs)
        self._log = get_logger("converters.word")
        self._forced_engine = engine
        self._selected_engine: Optional[str] = None

    # ------------------------------------------------------------------ #
    # 主入口
    # ------------------------------------------------------------------ #
    def convert(self, source: PathLike, target: PathLike, **kwargs: object) -> Path:
        """根据扩展名把 ``source`` 转换为 ``target``，返回实际写入路径。

        Args:
            source: 源文件路径（.docx / .doc / .pdf）。
            target: 目标文件路径。

        Returns:
            实际写入的目标 ``Path``。

        Raises:
            ConversionError: 不支持的扩展名组合、源文件不存在、依赖缺失
                或渲染进程失败。
        """
        src, dst = self._resolve_paths(source, target)
        src_ext, dst_ext = src.suffix.lower(), dst.suffix.lower()
        self._check_pair_supported(self.supported_pairs, src_ext, dst_ext)

        ensure_dir(dst.parent)
        final_dst = unique_output_path(dst)

        self._log.info(
            "开始转换: %s -> %s (engine=%s)",
            src.name,
            final_dst.name,
            self._forced_engine or "auto",
        )

        try:
            # ``on_progress`` 是可选进度回调 ``(processed, total) -> None``。
            # 当前只有 PDF -> DOCX 真正实现，Word -> PDF 直接忽略（不传）。
            on_progress = kwargs.pop("on_progress", None)
            if dst_ext == ".pdf":
                # Word -> PDF 没有页级概念，_to_pdf 不接受 on_progress
                result = self._to_pdf(src, final_dst, **kwargs)
            elif dst_ext == ".docx":
                result = self._pdf_to_docx(src, final_dst, on_progress=on_progress)
            else:  # 防御性代码：_check_pair_supported 已经过滤
                raise ConversionError(f"不支持的目标扩展名: {dst_ext}")
        except ConversionError:
            raise
        except Exception as exc:  # noqa: BLE001 - 包装为统一异常类型
            self._log.exception("Word 转换失败: %s", src)
            raise ConversionError(f"Word 转换失败 ({src.name} -> {final_dst.name}): {exc}") from exc

        self._log.info("转换完成: %s", result)
        return result

    # ------------------------------------------------------------------ #
    # Word -> PDF
    # ------------------------------------------------------------------ #
    def _to_pdf(self, src: Path, dst: Path) -> Path:
        """Word -> PDF：根据引擎选择调用对应实现。"""
        engine = self._forced_engine or self._select_engine()
        self._selected_engine = engine
        self._log.info("Word -> PDF 选用引擎: %s", engine)

        if engine == "docx2pdf":
            return self._to_pdf_via_docx2pdf(src, dst)
        if engine == "libreoffice":
            return self._to_pdf_via_libreoffice(src, dst)
        if engine == "pywin32":
            return self._to_pdf_via_pywin32(src, dst)
        raise ConversionError(f"未知的 Word->PDF 引擎: {engine}")

    def _to_pdf_via_docx2pdf(self, src: Path, dst: Path) -> Path:
        """通过 :mod:`docx2pdf` 调用本机 MS Word 完成转换。

        :mod:`docx2pdf` 没有输出文件名参数，只能指定输出目录并保留源
        文件名，因此转换完成后再 ``shutil.move`` 到目标位置。
        """
        if _docx2pdf_convert is None:  # pragma: no cover - 由 import 结果决定
            raise ConversionError(
                "docx2pdf 不可用：请 pip install docx2pdf，"
                "并在 Windows / macOS 上安装 Microsoft Word。"
            )

        # docx2pdf 在 Windows 上需要 COM 线程隔离，这里仅做单次同步调用。
        _docx2pdf_convert(str(src), str(dst.parent))

        produced = dst.parent / (src.stem + ".pdf")
        if not produced.exists():
            raise ConversionError(f"docx2pdf 未生成预期文件: {produced}")
        if produced != dst:
            shutil.move(str(produced), str(dst))
        return dst

    def _to_pdf_via_libreoffice(self, src: Path, dst: Path) -> Path:
        """通过 ``soffice --headless --convert-to pdf`` 完成转换。"""
        # 先查 PATH；找不到就回退到 office_status()（含 Windows 上的
        # Program Files 常见安装路径），避免 PATH 缺失导致误判。
        soffice = (
            shutil.which("soffice")
            or shutil.which("libreoffice")
            or office_status().soffice_path
        )
        if not soffice:
            raise ConversionError(
                "未找到 soffice/libreoffice 可执行文件，无法完成 Word->PDF。"
            )

        cmd = [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(dst.parent),
            str(src),
        ]
        self._log.debug("执行 LibreOffice: %s", " ".join(cmd))

        try:
            result = subprocess.run(
                cmd,
                check=True,
                capture_output=True,
                timeout=_LIBREOFFICE_TIMEOUT,
            )
        except subprocess.TimeoutExpired as exc:
            raise ConversionError(
                f"LibreOffice 转换超时（>{_LIBREOFFICE_TIMEOUT}s）: {src}"
            ) from exc
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or b"").decode("utf-8", errors="replace")
            raise ConversionError(
                f"LibreOffice 转换失败 (rc={exc.returncode}): {stderr.strip()}"
            ) from exc

        produced = dst.parent / (src.stem + ".pdf")
        if not produced.exists():
            raise ConversionError(f"LibreOffice 未生成预期文件: {produced}")
        if produced != dst:
            shutil.move(str(produced), str(dst))
        return dst

    def _to_pdf_via_pywin32(self, src: Path, dst: Path) -> Path:
        """Windows 备选方案：直接通过 win32com 调用 MS Word。"""
        if not _HAS_WIN32:  # pragma: no cover - 平台相关
            raise ConversionError("pywin32 不可用，仅 Windows 支持此引擎。")

        word = _win32_client.gencache.EnsureDispatch("Word.Application")  # type: ignore[union-attr]
        try:
            word.Visible = False
            # 0 = wdOpenFormatAuto，按文件内容自动决定格式
            doc = word.Documents.Open(str(src))
            try:
                doc.SaveAs(str(dst), _WD_FORMAT_PDF)
            finally:
                doc.Close(False)
        finally:
            word.Quit()
        if not dst.exists():
            raise ConversionError(f"pywin32 转换后未生成文件: {dst}")
        return dst

    # ------------------------------------------------------------------ #
    # PDF -> DOCX
    # ------------------------------------------------------------------ #
    def _pdf_to_docx(
        self,
        src: Path,
        dst: Path,
        on_progress: Optional[Callable[[int, int], None]] = None,
    ) -> Path:
        """从 PDF 抽取文本重新生成 Word 文档。

        实现要点：

        - 每个 PDF 页面对应一组 Word 段落；
        - 页面之间插入分页符；
        - 文本按原始行拆分（``splitlines``）逐行写入；
        - **会丢失字体、颜色、表格、列、页眉页脚、矢量图等版式信息**。
          适合对纯文字内容做"可编辑化"。

        Raises:
            ConversionError: 依赖 ``pdfplumber`` 或 ``python-docx`` 缺失。
        """
        if pdfplumber is None or _DocxDocument is None:  # pragma: no cover
            missing = []
            if pdfplumber is None:
                missing.append("pdfplumber")
            if _DocxDocument is None:
                missing.append("python-docx")
            raise ConversionError(
                f"PDF -> DOCX 缺少依赖: {', '.join(missing)}。"
                "请执行: pip install pdfplumber python-docx"
            )

        # 匹配纯数字页码行（含可选连字符装饰，如 "1" / "- 12 -" / "— 12 —"）。
        # pdfplumber 会把页眉页脚的页码当成普通文本行抓出来，
        # 必须剥离，否则会导致"页码单独成页"的问题。
        _PAGE_NUMBER_LINE = re.compile(r"^\s*[\-—–]?\s*\d{1,4}[\-—–]?\s*$")

        # 匹配"章节标题"用于在章节前加分页符。
        # 中文数字编号 + 顿号或英文句点（兼顾"一、xxx"和"八. xxx"两种写法）。
        _CHAPTER_HEADING = re.compile(r"^[一二三四五六七八九十]+[、.]\s*\S")

        # 匹配列表项开头（用于段间分隔）。
        # 覆盖："1. xxx" / "1）xxx" / "1)、xxx" / "（1）xxx" / "(1) xxx" / "(1).xxx"
        #      "第一、xxx" / "第一章 xxx"
        _LIST_ITEM = re.compile(
            r"^\s*"
            r"(?:"
            r"\d+[.、)）]"  # 数字 + 任意列表符号
            r"|"
            r"[（(]\d+[.）)]?"  # (数字 + 可选点)
            r"|"
            r"第[一二三四五六七八九十]+[、.\s]"  # 第X、/第X.
            r")"
            r"\s*\S"
        )

        # 段尾标点：行末出现这些字符 → 段落真边界，后续行进入新段。
        # 包含闭括号 "）)】』"：典型如 "自我介绍（参考）" 后是新段。
        # 不含 ":" / "：" / "、" / "," / "，" 因为它们常作为引出列表
        # 或句中停顿，不当段尾。
        _SENT_END = ("。", "！", "？", ";", "；", "）", ")", "】", "』")  # 。！？；;)）】』

        # 短小标题特征：行末是 ":" 或 "：" 且行长度 <= 15 → 当段尾。
        # 长度 > 15 的 "：" 通常是引出列表，不当段尾。
        _SHORT_HEADING_END = (":", "：")  # :：

        # 目录行特征：长省略号（"....."）。若某页超过半数行符合此模式，
        # 整页不分页（否则每条目录行都会被误判为章节、每条前加分页符）。
        _TOC_DOTS = re.compile(r"\.{5,}")

        def _is_sent_end(line: str) -> bool:
            """行末是否为段尾标点。"""
            s = line.rstrip()
            return bool(s) and s[-1] in _SENT_END

        def _is_short_heading_end(line: str) -> bool:
            """行末是否为短小标题（"准备方法：" / "答案："）。"""
            s = line.rstrip()
            return bool(s) and len(s) <= 15 and s[-1] in _SHORT_HEADING_END

        doc = _DocxDocument()
        try:
            with pdfplumber.open(str(src)) as pdf:  # type: ignore[union-attr]
                page_count = len(pdf.pages)
                self._log.info("PDF 解析: 共 %d 页", page_count)

                # 初始进度：报告"打开 PDF，准备解析"
                if on_progress is not None:
                    try:
                        on_progress(0, page_count)
                    except Exception:  # noqa: BLE001 - 回调失败不应阻塞转换
                        self._log.warning("on_progress 回调异常 (initial)", exc_info=True)

                for page_idx, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    # 1) 空页：保留一个空段落，避免正文挤在一起
                    if not text.strip():
                        doc.add_paragraph("")
                        continue
                    raw_lines = [ln for ln in text.splitlines() if ln.strip()]
                    # 2) 目录页判定：>= 50% 行含长省略号 → 整页不分页，
                    #    但行间仍按段尾规则合并（目录每条仍是独立段）
                    is_toc = bool(raw_lines) and sum(
                        1 for ln in raw_lines if _TOC_DOTS.search(ln)
                    ) >= len(raw_lines) * 0.5
                    # 3) 按行剥页码
                    lines = [
                        ln for ln in raw_lines
                        if not _PAGE_NUMBER_LINE.match(ln)
                    ]
                    if not lines:
                        continue
                    # 4) 把行合并成段落。
                    #    段间分隔触发条件（满足任一即换段）：
                    #    a) 行匹配章节标题
                    #    b) 行匹配列表项开头
                    #    c) 上一行段尾标点收尾
                    #    d) 上一行是短小标题（行末 ":" 且短）
                    #    目录页跳过规则 b（避免误把目录项当列表项）。
                    buf: list[str] = []  # 当前累积的段

                    def _flush() -> None:
                        """把 buf 合并成一段写进 docx。"""
                        if not buf:
                            return
                        doc.add_paragraph("".join(buf))

                    for idx, line in enumerate(lines):
                        stripped = line.strip()
                        is_chapter = bool(_CHAPTER_HEADING.match(stripped))
                        is_list = (
                            not is_toc
                            and bool(_LIST_ITEM.match(stripped))
                        )
                        if is_chapter:
                            # 章节前加分页符 + 段间分隔
                            _flush()
                            buf = []
                            doc.add_page_break()
                            buf.append(line)
                            continue
                        if is_list:
                            # 列表项：独立成段
                            _flush()
                            buf = []
                            buf.append(line)
                            continue
                        # 续行：上一行如果段尾/短标题 → 先 flush 当前 buf
                        if buf and (
                            _is_sent_end(buf[-1])
                            or _is_short_heading_end(buf[-1])
                        ):
                            _flush()
                            buf = []
                        buf.append(line)
                    _flush()

                # 页级进度回调：每解析完一页就推一次
                # （包含空页 / 仅页码的页面）。
                if on_progress is not None:
                    try:
                        on_progress(page_idx + 1, page_count)
                    except Exception:  # noqa: BLE001
                        self._log.warning(
                            "on_progress 回调异常 (page %d/%d)",
                            page_idx + 1,
                            page_count,
                            exc_info=True,
                        )
        finally:
            doc.save(str(dst))

        if not dst.exists():
            raise ConversionError(f"PDF -> DOCX 未能写出文件: {dst}")
        return dst

    # ------------------------------------------------------------------ #
    # 引擎选择
    # ------------------------------------------------------------------ #
    def _select_engine(self) -> str:
        """根据平台和可用 Office 套件选择最佳 Word -> PDF 引擎。

        策略：

        1. 优先 docx2pdf：能导入且 ``office_status().has_office`` 为真时
           直接使用（Windows / macOS + MS Word）。
        2. 否则若本机有 LibreOffice，则回退到 libreoffice 引擎。
        3. 都没检测到时抛 :class:`ConversionError`。

        Returns:
            ``"docx2pdf"`` / ``"libreoffice"`` 之一。
        """
        status = office_status()
        if _docx2pdf_convert is not None and status.has_office:
            return "docx2pdf"
        if status.has_libreoffice:
            return "libreoffice"
        raise ConversionError(
            "Word 转 PDF 需要 Microsoft Office（Windows/macOS）或 LibreOffice，"
            "当前环境均未检测到。\n" + status.install_hint()
        )


# ---------------------------------------------------------------------- #
# 模块加载时自动注册到全局 Registry
# ---------------------------------------------------------------------- #
from ..core.registry import Registry  # noqa: E402  放在底部避免与 BaseConverter 循环 import

Registry.register(WordConverter())
