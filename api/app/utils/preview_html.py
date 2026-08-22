
"""产物在线预览的 HTML 渲染工具。

把 XLSX / DOCX 产物转换为内嵌 HTML（供前端 iframe 预览），
PDF / 图片直接返回文件流（浏览器原生支持），无需转换。

设计：
- 所有用户内容一律 html.escape，防止 XSS。
- 大表保护：超过 MAX_ROWS / MAX_COLS 的行列截断并提示。
- .xls / .doc（旧二进制格式）无法用 openpyxl / python-docx 解析，
  返回占位提示 HTML。
"""

from __future__ import annotations

import html
from pathlib import Path
from typing import List

# 预览保护上限：避免超大文档生成超长 HTML
MAX_ROWS = 500
MAX_COLS = 50

# 包裹 HTML：统一页面样式
_PAGE_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<style>
  body { font-family: "Microsoft YaHei", "PingFang SC", sans-serif; margin: 16px; }
  h3 { margin: 16px 0 8px; }
  table { border-collapse: collapse; margin-bottom: 16px; }
  td, th { border: 1px solid #ccc; padding: 4px 8px; font-size: 13px; vertical-align: top; max-width: 320px; word-break: break-all; }
  th { background: #f5f7fa; }
  .notice { color: #909399; font-size: 13px; }
  .truncated { color: #e6a23c; font-size: 12px; margin: 4px 0 12px; }
</style>
</head>
<body>
%s
</body>
</html>"""


def _not_found_html(reason: str) -> str:
    """生成"无法预览"提示页。"""
    body = f'<p class="notice">{html.escape(reason)}</p>'
    return _PAGE_TEMPLATE % body


# ---------------------------------------------------------------------- #
# XLSX -> HTML
# ---------------------------------------------------------------------- #
def render_xlsx_html(path: Path) -> str:
    """把 .xlsx 渲染为 HTML 表格（保留合并单元格、Sheet 分节）。"""
    try:
        from openpyxl import load_workbook
    except Exception as exc:  # pragma: no cover
        return _not_found_html(f"缺少 openpyxl 依赖: {exc}")

    try:
        wb = load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        return _not_found_html(f"无法解析 Excel 文件: {exc}")

    try:
        parts: List[str] = []
        for ws in wb.worksheets:
            parts.append(f"<h3>{html.escape(ws.title)}</h3>")

            # 收集合并区域：anchor -> (rowspan, colspan)，以及全部覆盖坐标
            anchors: dict = {}
            covered: set = set()
            try:
                for rng in ws.merged_cells.ranges:
                    r_span = rng.max_row - rng.min_row + 1
                    c_span = rng.max_col - rng.min_col + 1
                    anchors[(rng.min_row, rng.min_col)] = (r_span, c_span)
                    for r in range(rng.min_row, rng.max_row + 1):
                        for c in range(rng.min_col, rng.max_col + 1):
                            covered.add((r, c))
            except Exception:
                pass  # read_only 模式下 merged_cells 可能不可用，忽略

            truncated = False
            parts.append("<table>")
            for i, row in enumerate(ws.iter_rows()):
                if i >= MAX_ROWS:
                    truncated = True
                    break
                tds: List[str] = []
                for cell in row:
                    if cell.column > MAX_COLS:
                        break
                    key = (cell.row, cell.column)
                    if key in anchors:
                        rs, cs = anchors[key]
                        val = cell.value
                        tds.append(
                            "<td" + (f" rowspan='{rs}' colspan='{cs}'" if (rs, cs) != (1, 1) else "")
                            + ">" + html.escape("" if val is None else str(val)) + "</td>"
                        )
                    elif key in covered:
                        continue  # 合并区非 anchor 格：跳过
                    else:
                        val = cell.value
                        tds.append(
                            "<td>" + html.escape("" if val is None else str(val)) + "</td>"
                        )
                if tds:
                    parts.append("<tr>" + "".join(tds) + "</tr>")
            parts.append("</table>")
            if truncated:
                parts.append(
                    '<p class="truncated">仅显示前 %d 行（共超过 %d 行），完整内容请下载查看。</p>'
                    % (MAX_ROWS, MAX_ROWS)
                )
        body = "".join(parts) if parts else '<p class="notice">工作簿为空</p>'
        return _PAGE_TEMPLATE % body
    finally:
        try:
            wb.close()
        except Exception:
            pass


# ---------------------------------------------------------------------- #
# DOCX -> HTML
# ---------------------------------------------------------------------- #
def render_docx_html(path: Path) -> str:
    """把 .docx 渲染为 HTML（段落 + 表格）。"""
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        return _not_found_html(f"缺少 python-docx 依赖: {exc}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        return _not_found_html(f"无法解析 Word 文件: {exc}")

    parts: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        parts.append(f"<p>{html.escape(text)}</p>")

    for table in doc.tables:
        parts.append("<table>")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(f"<td>{html.escape(cell.text.strip())}</td>")
            parts.append("<tr>" + "".join(cells) + "</tr>")
        parts.append("</table>")

    body = "".join(parts) if parts else '<p class="notice">文档没有可预览的文本内容</p>'
    return _PAGE_TEMPLATE % body


# ---------------------------------------------------------------------- #
# 统一入口
# ---------------------------------------------------------------------- #
def render_preview_html(path: Path) -> str:
    """按文件扩展名渲染预览 HTML。

    Returns:
        完整 HTML 文档字符串。

    Raises:
        ValueError: 扩展名不是 xlsx / docx（PDF、图片等应由调用方直接返回文件流）。
    """
    ext = path.suffix.lower()
    if ext == ".xlsx":
        return render_xlsx_html(path)
    if ext == ".docx":
        return render_docx_html(path)
    raise ValueError(f"不支持的预览 HTML 类型: {ext}")
